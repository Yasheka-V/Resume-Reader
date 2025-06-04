import os
import re
from docx import Document
import pandas as pd
from collections import Counter

def extract_years_of_experience(text):
    # Regex pattern to catch expressions like:
    # "X years", "X+ years", "over X years", "X years experience"
    pattern = r'(\d+)\s*\+?\s*(?:years|yrs)(?:\s*(?:of)?\s*experience)?'
    matches = re.findall(pattern, text, re.IGNORECASE)
    
  
    years_list = [int(num) for num in matches]
    if years_list:
        return max(years_list)
    else:
        return None
    
def extract_email(text):
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    matches = re.findall(email_pattern, text)
    return matches[0] if matches else None

def extract_phone(doc):
    phone_pattern = re.compile(
        r'(\+?\d{1,3}[-.\s]?)?'          # optional country code
        r'(\(?\d{3}\)?[-.\s]?)?'         # optional area code with or without parentheses
        r'\d{3}[-.\s]?\d{4,6}'           # main phone number
    )
    for para in doc.paragraphs:
        match = phone_pattern.search(para.text)
        if match:
            return match.group(0).strip()
    return None

def get_text_from_docx(path):
    doc = Document(path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def extract_name(doc):
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Take only the first line if there are multiple lines in the paragraph
            first_line = text.split('\n')[0].strip()
            # Skip if it looks like an email or contains digits (phone numbers or dates)
            if not re.match(r'[\w\.-]+@[\w\.-]+\.\w+', first_line) and not re.search(r'\d', first_line):
                return first_line
    return None

def calculate_match_percentage(resume_text, jd_text):
    resume_words = set(resume_text.split())
    jd_words = set(jd_text.split())
    common_words = resume_words.intersection(jd_words)
    if not jd_words:
        return 0
    match_percent = (len(common_words) / len(jd_words)) * 100
    return round(match_percent, 2)



Resume_Dir = r'C:\Users\Yasheka\AppData\Local\Programs\Python\Python313\Resumes'
JD_Dir = r'C:\Users\Yasheka\AppData\Local\Programs\Python\Python313\JD'
Output_File = r'C:\Users\Yasheka\AppData\Local\Programs\Python\Python313\Candidates_List.xlsx'
results = []


for filename in os.listdir(Resume_Dir):
    if filename.endswith('.docx'):
        filepath = os.path.join(Resume_Dir, filename)
        doc = Document(filepath)
        text = get_text_from_docx(filepath)
        name = extract_name(doc)
        email = extract_email(text)
        phone = extract_phone(doc)
        years = extract_years_of_experience(text)
        
        best_match = 0
        best_jd = None
        
        for JD in os.listdir(JD_Dir):
            if JD.endswith('.docx'):
                JD_Path = os.path.join(JD_Dir,JD)
                JD_doc = Document(JD_Path)
                JD_Text = get_text_from_docx(JD_Path)
                match = calculate_match_percentage(text, JD_Text)
                if match > best_match:
                    best_match = match
                    best_jd = os.path.basename(JD_Path)
        
        results.append({
            "Resume": filename,
            "Name": name if name else "Not found",
            "Email": email if email else "Not found",
            "Phone": phone if phone else "Not found",
            "Years of Experience": years if years is not None else "Not found",
            "Best Matching JD": best_jd,
            "Match Percentage": best_match
        })
        

                
       # Create DataFrame and save to Excel
df = pd.DataFrame(results)
df.to_excel(Output_File, index=False)

print(f"Results saved to {Output_File}") 
